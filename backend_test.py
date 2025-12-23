import requests
import sys
import os
from datetime import datetime
from pathlib import Path
import tempfile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PIL import Image
import io
from docx import Document
import openpyxl

class PDFMasterAPITester:
    def __init__(self, base_url="https://pdf-genius-24.preview.emergentagent.com"):
        self.base_url = base_url
        self.api_url = f"{base_url}/api"
        self.tests_run = 0
        self.tests_passed = 0
        self.failed_tests = []
        self.temp_files = []

    def create_test_pdf(self, filename="test.pdf", pages=2):
        """Create a simple test PDF file"""
        temp_path = Path(tempfile.gettempdir()) / filename
        c = canvas.Canvas(str(temp_path), pagesize=letter)
        
        for i in range(pages):
            c.drawString(100, 750 - (i * 50), f"This is page {i + 1}")
            c.drawString(100, 700 - (i * 50), "Sample PDF content for testing")
            if i < pages - 1:
                c.showPage()
        
        c.save()
        self.temp_files.append(temp_path)
        return temp_path

    def create_test_image(self, filename="test.jpg", format="JPEG"):
        """Create a simple test image file"""
        temp_path = Path(tempfile.gettempdir()) / filename
        img = Image.new('RGB', (200, 200), color='red')
        img.save(temp_path, format)
        self.temp_files.append(temp_path)
        return temp_path

    def create_test_word(self, filename="test.docx"):
        """Create a simple test Word document"""
        temp_path = Path(tempfile.gettempdir()) / filename
        doc = Document()
        doc.add_paragraph("This is a test Word document.")
        doc.add_paragraph("It contains sample content for testing.")
        doc.save(temp_path)
        self.temp_files.append(temp_path)
        return temp_path

    def create_test_excel(self, filename="test.xlsx"):
        """Create a simple test Excel file"""
        temp_path = Path(tempfile.gettempdir()) / filename
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = 'Test Data'
        ws['A2'] = 'Sample Excel Content'
        ws['B1'] = 'Column 2'
        ws['B2'] = 'More Data'
        wb.save(temp_path)
        self.temp_files.append(temp_path)
        return temp_path

    def run_test(self, name, method, endpoint, files=None, data=None, expected_status=200, response_type='file'):
        """Run a single API test"""
        url = f"{self.api_url}/{endpoint}"
        
        self.tests_run += 1
        print(f"\n🔍 Testing {name}...")
        
        try:
            if method == 'GET':
                response = requests.get(url, timeout=30)
            elif method == 'POST':
                if files:
                    response = requests.post(url, files=files, data=data, timeout=60)
                else:
                    response = requests.post(url, json=data, timeout=60)

            success = response.status_code == expected_status
            if success:
                self.tests_passed += 1
                print(f"✅ Passed - Status: {response.status_code}")
                
                # Check response content for file endpoints
                if response_type == 'file' and response.content:
                    print(f"   📄 Response size: {len(response.content)} bytes")
                elif response_type == 'json':
                    print(f"   📄 Response: {response.json()}")
                    
                return True, response
            else:
                print(f"❌ Failed - Expected {expected_status}, got {response.status_code}")
                if response.text:
                    print(f"   Error: {response.text[:200]}")
                self.failed_tests.append({
                    'name': name,
                    'endpoint': endpoint,
                    'expected': expected_status,
                    'actual': response.status_code,
                    'error': response.text[:200] if response.text else 'No error message'
                })
                return False, response

        except Exception as e:
            print(f"❌ Failed - Error: {str(e)}")
            self.failed_tests.append({
                'name': name,
                'endpoint': endpoint,
                'expected': expected_status,
                'actual': 'Exception',
                'error': str(e)
            })
            return False, None

    def test_root_endpoint(self):
        """Test the root API endpoint"""
        return self.run_test("Root API", "GET", "", expected_status=200, response_type='json')

    def test_merge_pdf(self):
        """Test PDF merge functionality"""
        pdf1 = self.create_test_pdf("merge1.pdf", 1)
        pdf2 = self.create_test_pdf("merge2.pdf", 1)
        
        files = [
            ('files', ('merge1.pdf', open(pdf1, 'rb'), 'application/pdf')),
            ('files', ('merge2.pdf', open(pdf2, 'rb'), 'application/pdf'))
        ]
        
        success, response = self.run_test("Merge PDF", "POST", "merge", files=files)
        
        # Close files
        for _, file_tuple in files:
            file_tuple[1].close()
            
        return success

    def test_split_pdf(self):
        """Test PDF split functionality"""
        pdf_file = self.create_test_pdf("split_test.pdf", 3)
        
        files = [('file', ('split_test.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        data = {'pages': '1-2'}
        
        success, response = self.run_test("Split PDF", "POST", "split", files=files, data=data)
        
        files[0][1][1].close()
        return success

    def test_compress_pdf(self):
        """Test PDF compression"""
        pdf_file = self.create_test_pdf("compress_test.pdf", 2)
        
        files = [('file', ('compress_test.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        
        success, response = self.run_test("Compress PDF", "POST", "compress", files=files)
        
        files[0][1][1].close()
        return success

    def test_rotate_pdf(self):
        """Test PDF rotation"""
        pdf_file = self.create_test_pdf("rotate_test.pdf", 1)
        
        files = [('file', ('rotate_test.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        data = {'angle': '90'}
        
        success, response = self.run_test("Rotate PDF", "POST", "rotate", files=files, data=data)
        
        files[0][1][1].close()
        return success

    def test_pdf_to_jpg(self):
        """Test PDF to JPG conversion"""
        pdf_file = self.create_test_pdf("pdf_to_jpg.pdf", 1)
        
        files = [('file', ('pdf_to_jpg.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        
        success, response = self.run_test("PDF to JPG", "POST", "pdf-to-jpg", files=files)
        
        files[0][1][1].close()
        return success

    def test_pdf_to_png(self):
        """Test PDF to PNG conversion"""
        pdf_file = self.create_test_pdf("pdf_to_png.pdf", 1)
        
        files = [('file', ('pdf_to_png.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        
        success, response = self.run_test("PDF to PNG", "POST", "pdf-to-png", files=files)
        
        files[0][1][1].close()
        return success

    def test_jpg_to_pdf(self):
        """Test JPG to PDF conversion"""
        jpg_file = self.create_test_image("test.jpg", "JPEG")
        
        files = [('file', ('test.jpg', open(jpg_file, 'rb'), 'image/jpeg'))]
        
        success, response = self.run_test("JPG to PDF", "POST", "jpg-to-pdf", files=files)
        
        files[0][1][1].close()
        return success

    def test_png_to_pdf(self):
        """Test PNG to PDF conversion"""
        png_file = self.create_test_image("test.png", "PNG")
        
        files = [('file', ('test.png', open(png_file, 'rb'), 'image/png'))]
        
        success, response = self.run_test("PNG to PDF", "POST", "png-to-pdf", files=files)
        
        files[0][1][1].close()
        return success

    def test_pdf_to_word(self):
        """Test PDF to Word conversion"""
        pdf_file = self.create_test_pdf("pdf_to_word.pdf", 1)
        
        files = [('file', ('pdf_to_word.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        
        success, response = self.run_test("PDF to Word", "POST", "pdf-to-word", files=files)
        
        files[0][1][1].close()
        return success

    def test_word_to_pdf(self):
        """Test Word to PDF conversion"""
        word_file = self.create_test_word("test.docx")
        
        files = [('file', ('test.docx', open(word_file, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))]
        
        success, response = self.run_test("Word to PDF", "POST", "word-to-pdf", files=files)
        
        files[0][1][1].close()
        return success

    def test_excel_to_pdf(self):
        """Test Excel to PDF conversion"""
        excel_file = self.create_test_excel("test.xlsx")
        
        files = [('file', ('test.xlsx', open(excel_file, 'rb'), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'))]
        
        success, response = self.run_test("Excel to PDF", "POST", "excel-to-pdf", files=files)
        
        files[0][1][1].close()
        return success

    def test_pdf_to_excel(self):
        """Test PDF to Excel conversion"""
        pdf_file = self.create_test_pdf("pdf_to_excel.pdf", 1)
        
        files = [('file', ('pdf_to_excel.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        
        success, response = self.run_test("PDF to Excel", "POST", "pdf-to-excel", files=files)
        
        files[0][1][1].close()
        return success

    def test_ocr_pdf(self):
        """Test OCR functionality"""
        pdf_file = self.create_test_pdf("ocr_test.pdf", 1)
        
        files = [('file', ('ocr_test.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        
        success, response = self.run_test("OCR PDF", "POST", "ocr", files=files, response_type='json')
        
        files[0][1][1].close()
        return success

    def test_watermark_pdf(self):
        """Test PDF watermarking"""
        pdf_file = self.create_test_pdf("watermark_test.pdf", 1)
        
        files = [('file', ('watermark_test.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        data = {'text': 'CONFIDENTIAL'}
        
        success, response = self.run_test("Watermark PDF", "POST", "watermark", files=files, data=data)
        
        files[0][1][1].close()
        return success

    def test_protect_pdf(self):
        """Test PDF password protection"""
        pdf_file = self.create_test_pdf("protect_test.pdf", 1)
        
        files = [('file', ('protect_test.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        data = {'password': 'testpass123'}
        
        success, response = self.run_test("Protect PDF", "POST", "protect", files=files, data=data)
        
        files[0][1][1].close()
        return success

    def test_unlock_pdf(self):
        """Test PDF password removal - this will fail without a protected PDF"""
        pdf_file = self.create_test_pdf("unlock_test.pdf", 1)
        
        files = [('file', ('unlock_test.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        data = {'password': 'wrongpass'}
        
        # This should fail since we're using an unprotected PDF
        success, response = self.run_test("Unlock PDF", "POST", "unlock", files=files, data=data, expected_status=500)
        
        files[0][1][1].close()
        return success

    def test_sign_pdf(self):
        """Test PDF signing"""
        pdf_file = self.create_test_pdf("sign_test.pdf", 1)
        
        files = [('file', ('sign_test.pdf', open(pdf_file, 'rb'), 'application/pdf'))]
        data = {'signature_text': 'John Doe'}
        
        success, response = self.run_test("Sign PDF", "POST", "sign", files=files, data=data)
        
        files[0][1][1].close()
        return success

    def cleanup(self):
        """Clean up temporary files"""
        for temp_file in self.temp_files:
            try:
                if temp_file.exists():
                    temp_file.unlink()
            except Exception as e:
                print(f"Warning: Could not delete {temp_file}: {e}")

    def run_all_tests(self):
        """Run all API tests"""
        print("🚀 Starting PDF Master API Tests...")
        print(f"📡 Testing against: {self.base_url}")
        
        # Test all endpoints
        test_methods = [
            self.test_root_endpoint,
            self.test_merge_pdf,
            self.test_split_pdf,
            self.test_compress_pdf,
            self.test_rotate_pdf,
            self.test_pdf_to_jpg,
            self.test_pdf_to_png,
            self.test_jpg_to_pdf,
            self.test_png_to_pdf,
            self.test_pdf_to_word,
            self.test_word_to_pdf,
            self.test_excel_to_pdf,
            self.test_pdf_to_excel,
            self.test_ocr_pdf,
            self.test_watermark_pdf,
            self.test_protect_pdf,
            self.test_unlock_pdf,
            self.test_sign_pdf
        ]
        
        for test_method in test_methods:
            try:
                test_method()
            except Exception as e:
                print(f"❌ Test {test_method.__name__} crashed: {e}")
                self.failed_tests.append({
                    'name': test_method.__name__,
                    'endpoint': 'unknown',
                    'expected': 'success',
                    'actual': 'crash',
                    'error': str(e)
                })
        
        # Print results
        print(f"\n📊 Test Results:")
        print(f"✅ Passed: {self.tests_passed}/{self.tests_run}")
        print(f"❌ Failed: {len(self.failed_tests)}")
        
        if self.failed_tests:
            print(f"\n🔍 Failed Tests:")
            for test in self.failed_tests:
                print(f"  - {test['name']}: {test['error'][:100]}")
        
        # Cleanup
        self.cleanup()
        
        return self.tests_passed == self.tests_run

def main():
    tester = PDFMasterAPITester()
    success = tester.run_all_tests()
    return 0 if success else 1

if __name__ == "__main__":
    sys.exit(main())